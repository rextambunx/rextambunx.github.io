from fastapi import FastAPI, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
from docxtpl import DocxTemplate
import json
import os
import uuid
import uvicorn
from fastapi import FastAPI, Query, HTTPException
from fastapi.responses import StreamingResponse
from docx import Document
import io
import json

app = FastAPI()

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # หรือระบุ frontend origin เช่น "http://localhost:3000"
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

TEMPLATE_FILE = "templates/template.docx"
ANSWERS_FILE = "data/saved_answers.json"
OUTPUT_FOLDER = "generated_docs"

os.makedirs(OUTPUT_FOLDER, exist_ok=True)

@app.get("/list-answers")
def list_answers():
    with open("data/saved_answers.json", "r", encoding="utf-8") as f:
        all_answers_list = json.load(f)

    keys = []
    for answer_obj in all_answers_list:
        keys.extend(answer_obj.keys())

    return {"keys": keys}


@app.get("/generate-docx")
def generate_docx(answer_key: str = Query(...)):
    try:
        with open("data/saved_answers.json", "r", encoding="utf-8") as f:
            all_answers_list = json.load(f)

        # แปลง list ของ object ให้เป็น dict ปกติ
        all_answers = {}
        for answer_obj in all_answers_list:
            all_answers.update(answer_obj)  # รวมเข้า dict

        answer_data = all_answers.get(answer_key)
        if not answer_data:
            raise HTTPException(status_code=404, detail="Answer key not found")

        # ---- เหมือนเดิม ----
        doc = Document()
        doc.add_heading("ข้อตกลงผู้ถือหุ้น", level=1)

        doc.add_paragraph(f"บริษัท: {answer_data.get("q29: คำถาม 2.5.1: บริษัทชื่ออะไร?", '')}")
        doc.add_paragraph(f"สถานที่ลงนาม: {answer_data.get("q3: คำถาม 1.1.1: ข้อตกลงจะลงนามที่ไหน?", '')}")
        doc.add_paragraph(f"วันที่ลงนาม: {answer_data.get("q5: คำถาม 1.1.2: ข้อตกลงจะลงนามในวันไหน?", '')}")

        doc.add_heading("ผู้ถือหุ้น:", level=2)
        doc.add_paragraph(f"1. {answer_data.get( "q6: คำถาม 1.1.3: ตามกฎหมายบริษัทไทย บริษัทจำกัดจำเป็นต้องมีผู้ถือหุ้นอย่างน้อยสามคน ผู้ถือหุ้นคนแรกชื่ออะไร?", '')}")
        doc.add_paragraph(f"2. {answer_data.get( "q7: คำถาม 1.1.3.1: ผู้ถือหุ้นคนที่สองชื่ออะไร?", '')}")
        doc.add_paragraph(f"3. {answer_data.get("q8: คำถาม 1.1.3.1.1: ผู้ถือหุ้นคนที่สามชื่ออะไร?", '')}")

        doc.add_paragraph(f"\nวันที่มีผล: {answer_data.get('q16:  คำถาม 2.1.1.1.1: ข้อตกลงนี้จะมีผลบังคับใช้ในวันไหน?', '')}")
        doc.add_paragraph(f"ลงชื่อ: {answer_data.get("q28: คำถาม 2.4.1 : [ชื่อ, นามสกุล, เลขทะเบียน]", '')}")
        doc.add_paragraph(f"บริษัทจะยุบเลิกเมื่อไหร่: {answer_data.get("q32: คำถาม 2.5.3.1: บริษัทจะยุบเลิกเมื่อไหร่?", '')}")

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        filename = f"{answer_key.replace(' ', '_')}.docx"
        return StreamingResponse(
            buffer,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename={filename}"}
        )

    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


# รันเซิร์ฟเวอร์
if __name__ == "__main__":
    uvicorn.run("Time:app", host="0.0.0.0", port=8001, reload=True)